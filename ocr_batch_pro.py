#!/usr/bin/env python3
"""
OCR PDF Converter - Batch Processing

Batch-converts scanned PDFs to Word (.docx) or plain text (.txt) using OCR.
Runs on macOS, Windows and Linux: file pickers come from Tk (bundled with
Python), and the two external binaries - Tesseract and Poppler - are located
per platform at start-up rather than assumed to be on PATH.
"""

import pdf2image
import pytesseract
import cv2
import numpy as np
from PIL import Image
import os
import shutil
import subprocess
import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

IS_WINDOWS = sys.platform.startswith("win")
IS_MAC = sys.platform == "darwin"


# --------------------------------------------------------------------------- #
# Locating the two external binaries.
#
# On macOS and Linux a package manager puts these on PATH. On Windows the
# installers do not touch PATH by default, so the usual install locations are
# checked explicitly - otherwise the failure is an opaque error from deep
# inside pytesseract or pdf2image rather than something a person can act on.
# --------------------------------------------------------------------------- #

# The launcher can download portable copies of these into ./tools rather than
# installing anything system-wide, so that folder is searched before anything
# else - a tool sitting next to the script is the one this app was set up with.
APP_DIR = Path(__file__).resolve().parent
TOOLS_DIR = APP_DIR / "tools"


def _in_tools(exe_name):
    """First matching executable anywhere under ./tools, or None."""
    if not TOOLS_DIR.is_dir():
        return None
    for found in sorted(TOOLS_DIR.rglob(exe_name)):
        if found.is_file():
            return str(found)
    return None


TESSERACT_CANDIDATES = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    "/opt/homebrew/bin/tesseract",
    "/usr/local/bin/tesseract",
    "/usr/bin/tesseract",
]

# The Windows poppler build unzips to a version-stamped folder such as
# "poppler-26.02.0", so the Windows locations are globbed rather than listed.
POPPLER_BIN_GLOBS = [
    r"C:\Program Files\poppler*\Library\bin",
    r"C:\Program Files\poppler*\bin",
    r"C:\poppler*\Library\bin",
    r"C:\poppler*\bin",
]

POPPLER_BIN_CANDIDATES = [
    "/opt/homebrew/bin",
    "/usr/local/bin",
    "/usr/bin",
]

POPPLER_PATH = None   # passed to pdf2image when the binaries are not on PATH


def _find_tesseract():
    exe = "tesseract.exe" if IS_WINDOWS else "tesseract"
    return (_in_tools(exe)
            or shutil.which("tesseract")
            or next((c for c in TESSERACT_CANDIDATES if os.path.isfile(c)), None))


def _find_poppler_bin():
    """Return the folder holding pdftoppm, None if it is already on PATH, or
    the string "MISSING" if it cannot be found at all."""
    exe_name = "pdftoppm.exe" if IS_WINDOWS else "pdftoppm"
    local = _in_tools(exe_name)
    if local:
        return os.path.dirname(local)
    if shutil.which("pdftoppm") or shutil.which("pdftoppm.exe"):
        return None
    exe = "pdftoppm.exe" if IS_WINDOWS else "pdftoppm"
    candidates = list(POPPLER_BIN_CANDIDATES)
    if IS_WINDOWS:
        import glob
        for pattern in POPPLER_BIN_GLOBS:
            candidates = sorted(glob.glob(pattern), reverse=True) + candidates
    for d in candidates:
        if os.path.isfile(os.path.join(d, exe)):
            return d
    return "MISSING"


def check_dependencies():
    """Fail early, with an instruction the person can actually follow."""
    global POPPLER_PATH
    problems = []

    tess = _find_tesseract()
    if tess:
        pytesseract.pytesseract.tesseract_cmd = tess
    else:
        if IS_WINDOWS:
            problems.append(
                "Tesseract OCR was not found.\n"
                "    Install it:  winget install -e --id UB-Mannheim.TesseractOCR\n"
                "    (or download the installer from "
                "https://github.com/UB-Mannheim/tesseract/wiki)")
        elif IS_MAC:
            problems.append("Tesseract OCR was not found.\n"
                            "    Install it:  brew install tesseract")
        else:
            problems.append("Tesseract OCR was not found.\n"
                            "    Install it:  sudo apt install tesseract-ocr")

    poppler = _find_poppler_bin()
    if poppler == "MISSING":
        if IS_WINDOWS:
            problems.append(
                "Poppler was not found (needed to turn PDF pages into images).\n"
                "    Download the latest release from "
                "https://github.com/oschwartz10612/poppler-windows/releases\n"
                "    and unzip it to C:\\Program Files\\poppler")
        elif IS_MAC:
            problems.append("Poppler was not found.\n"
                            "    Install it:  brew install poppler")
        else:
            problems.append("Poppler was not found.\n"
                            "    Install it:  sudo apt install poppler-utils")
    else:
        POPPLER_PATH = poppler   # None means "already on PATH"

    if problems:
        print("\nMissing requirements:\n")
        for p in problems:
            print("  - " + p + "\n")
        return False
    return True


class OCRConverter:
    """Handles PDF → OCR → Word/TXT/Pages conversion with batch support."""

    def __init__(self, output_format, preprocess, denoise_strength, threshold_value):
        self.output_format = output_format
        self.preprocess = preprocess
        self.denoise_strength = denoise_strength
        self.threshold_value = threshold_value

    def preprocess_image(self, image):
        """Apply denoise and threshold to image."""
        try:
            cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)

            if self.denoise_strength > 0:
                denoised = cv2.fastNlMeansDenoising(
                    gray,
                    h=self.denoise_strength,
                    templateWindowSize=7,
                    searchWindowSize=21
                )
            else:
                denoised = gray

            if self.threshold_value > 0:
                thresholded = cv2.adaptiveThreshold(
                    denoised,
                    255,
                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY,
                    blockSize=self.threshold_value,
                    C=2
                )
            else:
                thresholded = denoised

            return Image.fromarray(thresholded)
        except Exception as e:
            print(f"⚠️  Preprocessing error: {str(e)}")
            return image

    def extract_text(self, pdf_path):
        """Extract text from PDF."""
        try:
            print(f"  📄 Loading: {Path(pdf_path).name}")
            kw = {'dpi': 300, 'fmt': 'ppm'}
            if POPPLER_PATH:
                kw['poppler_path'] = POPPLER_PATH
            images = pdf2image.convert_from_path(pdf_path, **kw)
            total_pages = len(images)
            print(f"  ✓ {total_pages} pages")

            extracted_pages = []

            for page_num in range(total_pages):
                try:
                    image = images[page_num]
                    if self.preprocess:
                        image = self.preprocess_image(image)
                    text = pytesseract.image_to_string(image)
                    extracted_pages.append({
                        'page_num': page_num + 1,
                        'text': text.strip() if text else "[No text detected]"
                    })
                except Exception as e:
                    extracted_pages.append({
                        'page_num': page_num + 1,
                        'text': f"[OCR failed]"
                    })

            return extracted_pages
        except Exception as e:
            print(f"  ❌ Error: {str(e)}")
            return None

    def save_as_docx(self, pdf_path, output_dir, extracted_pages):
        """Save as Word document."""
        try:
            doc = Document()
            pdf_name = Path(pdf_path).stem

            title = doc.add_heading('OCR Conversion Report', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            meta = doc.add_paragraph(f"Source: {pdf_name}\nPages: {len(extracted_pages)}")
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            for page_data in extracted_pages:
                doc.add_heading(f'Page {page_data["page_num"]}', level=2)
                doc.add_paragraph(page_data['text'])

            output_file = os.path.join(output_dir, f"{pdf_name}_OCR.docx")
            doc.save(output_file)
            return output_file
        except Exception as e:
            print(f"  ❌ Error saving DOCX: {str(e)}")
            return None

    def save_as_txt(self, pdf_path, output_dir, extracted_pages):
        """Save as plain text."""
        try:
            pdf_name = Path(pdf_path).stem
            output_file = os.path.join(output_dir, f"{pdf_name}_OCR.txt")

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(f"OCR Conversion: {pdf_name}\n")
                f.write(f"Pages: {len(extracted_pages)}\n")
                f.write("=" * 80 + "\n\n")

                for page_data in extracted_pages:
                    f.write(f"--- Page {page_data['page_num']} ---\n")
                    f.write(page_data['text'])
                    f.write("\n\n")

            return output_file
        except Exception as e:
            print(f"  ❌ Error saving TXT: {str(e)}")
            return None


    def convert_pdf(self, pdf_path, output_dir):
        """Convert a single PDF."""
        extracted_pages = self.extract_text(pdf_path)
        if not extracted_pages:
            return None

        if self.output_format == 'Word (.docx)':
            return self.save_as_docx(pdf_path, output_dir, extracted_pages)
        elif self.output_format == 'Plain Text (.txt)':
            return self.save_as_txt(pdf_path, output_dir, extracted_pages)


# --------------------------------------------------------------------------- #
# File pickers.
#
# Tk ships with Python on every platform, so one dialog implementation works
# everywhere. If Tk is genuinely unavailable (a stripped Linux Python, a
# headless session) the prompts fall back to typed paths rather than failing.
# --------------------------------------------------------------------------- #

def _tk_root():
    """A hidden Tk root, or None if Tk cannot start here."""
    try:
        import tkinter as tk
        root = tk.Tk()
        root.withdraw()
        root.update()
        try:
            root.attributes("-topmost", True)   # dialogs above the console
        except Exception:
            pass
        return root
    except Exception:
        return None


def pick_pdfs():
    """Pick one or more PDFs. Returns a list of paths (empty if cancelled)."""
    root = _tk_root()
    if root is not None:
        try:
            from tkinter import filedialog
            chosen = filedialog.askopenfilenames(
                title="Select PDF file(s)",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            )
            return list(chosen)
        finally:
            root.destroy()

    typed = input("  Path to a PDF (or a folder of PDFs): ").strip().strip('"\'')
    if not typed:
        return []
    path = Path(typed).expanduser()
    if path.is_dir():
        return [str(f) for f in sorted(path.glob("*.pdf"))]
    return [str(path)] if path.is_file() else []


def pick_output_folder():
    """Pick the output folder. Defaults to Downloads if cancelled."""
    default = str(Path.home() / "Downloads")
    root = _tk_root()
    if root is not None:
        try:
            from tkinter import filedialog
            chosen = filedialog.askdirectory(title="Select output folder")
            return chosen or default
        finally:
            root.destroy()

    typed = input(f"  Output folder (Enter for {default}): ").strip().strip('"\'')
    return str(Path(typed).expanduser()) if typed else default


def open_folder(path):
    """Reveal a folder in the platform's file manager."""
    try:
        if IS_WINDOWS:
            os.startfile(path)                              # noqa: S606
        elif IS_MAC:
            subprocess.run(["open", path], check=False)
        else:
            subprocess.run(["xdg-open", path], check=False)
    except Exception as e:
        print(f"  Could not open the folder automatically ({e}).")
        print(f"  It is here: {path}")


def main():
    """Main application."""
    print("\n" + "="*70)
    print("  OCR PDF Converter (Batch Processing)")
    print("="*70 + "\n")

    if not check_dependencies():
        print("=" * 70 + "\n")
        return

    # Pick PDFs
    print("📁 Select PDF file(s)...")
    pdf_files = pick_pdfs()
    if not pdf_files or pdf_files[0] == '':
        print("❌ No PDFs selected")
        return

    # Pick output folder
    print("📁 Select output folder...")
    output_dir = pick_output_folder()

    # Ask for format
    print("\n📋 Output format:")
    print("  1) Word (.docx)")
    print("  2) Plain Text (.txt)")
    choice = input("  Enter choice (1-2): ").strip()

    formats = {
        '1': 'Word (.docx)',
        '2': 'Plain Text (.txt)'
    }
    output_format = formats.get(choice, 'Word (.docx)')

    # Ask for preprocessing
    preprocess_choice = input("\n🔧 Enable preprocessing? (y/n, default y): ").strip().lower()
    preprocess = preprocess_choice != 'n'

    denoise = 10
    threshold = 11

    if preprocess:
        try:
            denoise_str = input("  Denoise strength (0-20, default 10): ").strip()
            denoise = int(denoise_str) if denoise_str else 10
        except:
            pass

        try:
            threshold_str = input("  Threshold block size (0-50, default 11): ").strip()
            threshold = int(threshold_str) if threshold_str else 11
        except:
            pass

    # Process files
    print("\n" + "="*70)
    converter = OCRConverter(output_format, preprocess, denoise, threshold)

    results = []
    for i, pdf_path in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}] Processing...")
        output_file = converter.convert_pdf(pdf_path, output_dir)
        if output_file:
            print(f"  ✅ Saved: {Path(output_file).name}")
            results.append(output_file)
        else:
            print(f"  ❌ Failed")

    # Summary
    print("\n" + "="*70)
    if results:
        print(f"✅ Completed: {len(results)}/{len(pdf_files)} files")
        print(f"📁 Output folder: {output_dir}\n")

        # Ask to open results
        open_choice = input("Open output folder? (y/n): ").strip().lower()
        if open_choice == 'y':
            open_folder(output_dir)
    else:
        print("❌ No files converted")

    print("="*70 + "\n")


if __name__ == '__main__':
    main()
